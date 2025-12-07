# ui/dashboard_view.py

import streamlit as st
import pandas as pd
import json
import asyncio
import datetime
import ast
import os
from collections import Counter
from services import llm_handler as llm
from services.database import DatabaseManager 
import plotly.express as px
from io import BytesIO
from docx import Document


EVIDENCE_DIR = "session_evidence"


def _create_docx(text_content):
    """
    將文字內容轉換為 Word 文件 (BytesIO)
    """
    doc = Document()
    doc.add_heading('DineSence AI Report', 0)
    
    # 簡單處理：按行寫入，保留段落感
    for line in text_content.split('\n'):
        doc.add_paragraph(line)
        
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# 輔助函式：圖片 Grid
def _render_evidence_grid(db_manager, session_id, event_type):
    """
    負責讀取並顯示圖片 Grid，不包含外層的 Expander。
    """
    evidence_df = db_manager.get_event_evidence(session_id, event_type)
    
    if evidence_df.empty:
        st.info("NO EVIDENCE FOUND.", icon="ℹ️")
        return

    # 限制顯示數量，避免一次載入太多當機
    cols = st.columns(4) 
    
    for i, row in evidence_df.iterrows():
        filename = os.path.basename(row['local_path'])
        path = os.path.join(EVIDENCE_DIR, filename)
        
        evidence_id = row['id']
        is_correct = row['human_corrected']
        
        # 決定圖片放在哪一欄
        col = cols[i % 4]
        
        if os.path.exists(path):
            with col:
                st.image(path, use_container_width=True)
                
                checkbox_key = f"img_feedback_{evidence_id}_{session_id}"
                
                def update_feedback(eid=evidence_id, key=checkbox_key):
                    new_state = st.session_state[key]
                    db_manager.update_evidence_feedback(eid, new_state)
                    if new_state:
                        st.toast(f"✅ ID {eid} CONFIRMED", icon="👍")

                st.checkbox(
                    f"#{evidence_id} CONFIRM", 
                    value=(is_correct == 1), 
                    key=checkbox_key,
                    on_change=update_feedback
                )
        else:
            with col:
                st.warning(f"MISSING {evidence_id}")

def _render_global_insights(client, db_manager, df_sessions, t):
    """
    [NEW] 總體數據洞察：跨 Session 的菜色情緒統計與 LLM 報告
    """
    st.info("此頁面統計範圍為上方「篩選器」所選定之時間段內的數據。")

    if df_sessions.empty:
        st.warning("⚠️ 目前選定的時間範圍內無 Session 資料。")
        return

    # 1. 跨 Session 資料聚合 (Data Aggregation)
    all_food_data = []
    
    # 遍歷篩選出的所有 Session
    for _, session_row in df_sessions.iterrows():
        sid = session_row['session_id_raw']
        
        # 撈取該 Session 的餐點證據
        df_evidence = db_manager.get_event_evidence(sid, "strong_emotion_plate")
        
        if df_evidence.empty: continue
            
        for _, row in df_evidence.iterrows():
            # 排除人工否決的資料
            if row['human_corrected'] == 0: continue
                
            food_name = row['food_label'] if row['food_label'] else "Unknown"
            
            # 解析情緒 (從檔名)
            # 檔名格式: 時間_情緒-分_...
            try:
                fname = os.path.basename(row['local_path'])
                parts = fname.split('_')
                emotion_tag = parts[2].split('-')[0] # 取出 "開心"
                
                all_food_data.append({
                    "session_id": sid,
                    "evidence_id": row['id'],
                    "food": food_name,
                    "emotion": emotion_tag,
                    "path": row['local_path'],
                    "timestamp": row['session_timestamp'] # 這裡可能是 session_id，需注意顯示格式
                })
            except:
                continue

    if not all_food_data:
        st.warning("⚠️ 在此時間範圍內，尚未偵測到任何有效的餐點情緒數據。")
        return

    df_analysis = pd.DataFrame(all_food_data)

    # 2. 統計數據準備 (給 LLM 用)
    # 格式: {'漢堡': {'開心': 5, '嫌棄': 1}, '薯條': ...}
    food_stats = {}
    for food in df_analysis['food'].unique():
        sub_df = df_analysis[df_analysis['food'] == food]
        counts = sub_df['emotion'].value_counts().to_dict()
        food_stats[food] = counts

    # ==========================================
    # 區塊 A: LLM 總體洞察報告
    # ==========================================
    with st.container(border=True):
        st.subheader("🤖 AI 營運洞察報告")
        st.markdown("讓 AI 為您分析本時段內，各項餐點的顧客情緒表現。")
        
        if st.button(t("btn_gen_insight_report"), type="primary", use_container_width=True):
            if not client:
                st.error("未設定 OpenAI API Key")
            else:
                with st.spinner("AI 正在分析大數據..."):
                    # 組建 Prompt
                    stats_str = json.dumps(food_stats, ensure_ascii=False, indent=2)
                    system_prompt = (
                        "你是一位專業的餐廳數據分析師。使用者會提供一份 JSON 數據，"
                        "內容是不同菜色對應的顧客情緒統計 (例如: 漢堡 -> 開心:5, 嫌棄:2)。\n"
                        "請根據數據生成一份繁體中文報告，包含：\n"
                        "1. 🏆 **明星菜色**：哪道菜的正面情緒(開心/驚艷)比例最高？\n"
                        "2. ⚠️ **改進建議**：哪道菜出現了負面情緒(嫌棄/失望/不滿)？可能原因？\n"
                        "3. 💡 **總結洞察**：整體菜單的表現評價。\n"
                        "請用專業、簡潔的條列式語氣回答。"
                    )
                    user_prompt = f"請分析以下餐點情緒數據：\n{stats_str}"

                    async def run_gpt():
                        try:
                            resp = await client.chat.completions.create(
                                model="gpt-4o",
                                messages=[
                                    {"role": "system", "content": system_prompt},
                                    {"role": "user", "content": user_prompt}
                                ],
                                temperature=0.7
                            )
                            return resp.choices[0].message.content
                        except Exception as e:
                            return f"Error: {e}"
                            
                    report_text = asyncio.run(run_gpt())
                    st.markdown("---")
                    st.markdown(report_text)

    st.divider()

    # ==========================================
    # 區塊 B: 單品項詳細分析
    # ==========================================
    c1, c2 = st.columns([1, 2])
    
    with c1:
        st.markdown("### 🍔 菜色細節查詢")
        food_list = sorted(list(food_stats.keys()))
        selected_food = st.selectbox("選擇要鑽研的菜色", food_list)
        
        # 顯示該菜色的基本數據
        if selected_food:
            stats = food_stats[selected_food]
            total = sum(stats.values())
            st.caption(f"共蒐集到 {total} 筆反應")
            st.json(stats)

    with c2:
        if selected_food:
            # 畫圖
            df_target = df_analysis[df_analysis['food'] == selected_food]
            emo_counts = df_target['emotion'].value_counts().reset_index()
            emo_counts.columns = ['Emotion', 'Count']
            
            fig = px.bar(
                emo_counts, x='Emotion', y='Count',
                title=f"「{selected_food}」情緒分佈圖",
                color='Emotion', text_auto=True,
                color_discrete_sequence=px.colors.qualitative.Pastel
            )
            st.plotly_chart(fig, use_container_width=True)

    # ==========================================
    # 區塊 C: 證據驗證與修正
    # ==========================================
    st.subheader(f"✅ 資料驗證 ({selected_food})")
    
    target_records = df_analysis[df_analysis['food'] == selected_food]
    
    # Grid 顯示
    cols = st.columns(4)
    for i, row in target_records.iterrows():
        col = cols[i % 4]
        with col:
            with st.container(border=True):
                # 顯示圖片
                if os.path.exists(row['path']):
                    st.image(row['path'], use_container_width=True)
                else:
                    st.warning("影像遺失")
                
                # 情緒標籤
                st.markdown(f"**{row['emotion']}**")
                
                # 勾選框
                chk_key = f"g_chk_{row['evidence_id']}"
                
                def update_cb(eid=row['evidence_id'], k=chk_key):
                    val = st.session_state[k]
                    db_manager.update_evidence_feedback(eid, val)
                    if not val: st.toast(f"已從統計中移除 (ID: {eid})")

                st.checkbox("確認無誤", value=True, key=chk_key, on_change=update_cb)

def _render_comparison_gallery(db_manager, session_id):
    """
    [NEW] 強烈情緒交叉比對畫廊
    邏輯：找出同一時間點的 Face 與 Plate 照片，並排顯示。
    """
    # 1. 撈出該 Session 所有強烈情緒相關的證據
    df_face = db_manager.get_event_evidence(session_id, "strong_emotion_face")
    df_plate = db_manager.get_event_evidence(session_id, "strong_emotion_plate")
    
    if df_face.empty and df_plate.empty:
        st.info("尚未偵測到強烈情緒事件 (Confidence > 40%)")
        return

    # 2. 進行配對 (Pairing)
    # 我們利用檔名中的時間戳記 (例如 "11月30日_12點01分05秒") 來配對
    pairs = {} 
    
    # 處理臉部照片
    for _, row in df_face.iterrows():
        path = row['local_path']
        filename = os.path.basename(path)
        
        # 修正後的解析邏輯
        parts = filename.split('_')
        # 取前兩段當作唯一的時間 Key (月日_時分秒)
        key = f"{parts[0]}_{parts[1]}"
        
        # 取得主要情緒名稱 (移除分數)
        raw_emo = parts[2] # "開心-98"
        emo_label = raw_emo.split('-')[0] # "開心"

        if key not in pairs: pairs[key] = {}
        pairs[key]['face'] = path
        pairs[key]['emotion'] = emo_label
        pairs[key]['time'] = parts[1] # 顯示時間

    # 處理餐盤照片
    for _, row in df_plate.iterrows():
        path = row['local_path']
        filename = os.path.basename(path)
        parts = filename.split('_')
        key = f"{parts[0]}_{parts[1]}"
        
        if key not in pairs: pairs[key] = {}
        pairs[key]['plate'] = path

    # 3. 渲染 UI (由新到舊排序)
    sorted_keys = sorted(pairs.keys(), reverse=True)
    
    for key in sorted_keys:
        item = pairs[key]
        face_path = item.get('face')
        plate_path = item.get('plate')
        emotion_label = item.get('emotion', 'Unknown')
        time_label = item.get('time', '')

        # 卡片式佈局
        with st.container(border=True):
            # 標題列：顯示情緒與時間
            st.markdown(f"#### 🔥 {emotion_label} <span style='font-size:0.8em; color:gray'>({time_label})</span>", unsafe_allow_html=True)
            
            c1, c2 = st.columns(2)
            
            # 左邊：表情
            with c1:
                st.caption("👤 顧客表情")
                if face_path and os.path.exists(face_path):
                    st.image(face_path, use_container_width=True)
                else:
                    st.warning("影像遺失")
            
            # 右邊：餐盤
            with c2:
                st.caption("🍽️ 當下餐盤")
                if plate_path and os.path.exists(plate_path):
                    st.image(plate_path, use_container_width=True)
                else:
                    st.warning("影像遺失")
def _render_food_insights(db_manager, session_id):
    """
    [NEW] 餐點洞察模式：以食物為中心，統計顧客的情緒反應
    """
    # 1. 撈取該 Session 所有「強烈情緒的餐盤照」
    # 這些照片已經經過 LLM 辨識，帶有 food_label
    df_plate = db_manager.get_event_evidence(session_id, "strong_emotion_plate")
    
    if df_plate.empty:
        st.info("尚無 AI 辨識的餐點數據")
        return

    # 2. 資料前處理：解析檔名中的情緒，並過濾無效數據
    data_list = []
    
    for _, row in df_plate.iterrows():
        # 如果使用者已經手動取消勾選 (human_corrected=0)，就排除這筆資料
        if row['human_corrected'] == 0:
            continue
            
        food_name = row['food_label'] if row['food_label'] else "Unknown"
        path = row['local_path']
        evidence_id = row['id']
        
        # 從檔名解析情緒
        # 格式: {時間}_{情緒1-分}_{情緒2-分}_Plate.jpg
        # 範例: 12月01日_..._開心-98_驚艷-02_Plate.jpg
        try:
            filename = os.path.basename(path)
            parts = filename.split('_')
            
            # 取出第一高分的情緒
            e1_tag = parts[2] # "開心-98"
            emotion_label = e1_tag.split('-')[0] # "開心"
            
            # 為了顯示方便，我們也嘗試找對應的臉部照片
            # 只要把檔名結尾的 Plate.jpg 改成 Face.jpg 即可
            face_path = path.replace("_Plate.jpg", "_Face.jpg")
            
            data_list.append({
                "id": evidence_id,
                "food": food_name,
                "emotion": emotion_label,
                "plate_path": path,
                "face_path": face_path,
                "timestamp": parts[1]
            })
        except:
            continue

    if not data_list:
        st.warning("沒有有效的餐點數據 (可能都被取消勾選了)")
        return

    df_analysis = pd.DataFrame(data_list)

   # 3. UI 佈局
    all_foods = sorted(df_analysis['food'].unique().tolist())
    
    c1, c2 = st.columns([1, 3])
    with c1:
        st.markdown("### 🍔 選擇餐點")
        # ★★★ [修正] 加上 key 參數，綁定 session_id ★★★
        selected_food = st.selectbox(
            "請選擇要分析的菜色", 
            all_foods, 
            key=f"food_select_{session_id}" 
        )
    
    # 篩選出該食物的資料
    df_target = df_analysis[df_analysis['food'] == selected_food]
    
    with c2:
        # ==========================================
        # UI 區塊 B: 統計直方圖
        # ==========================================
        if not df_target.empty:
            # 統計各種情緒的出現次數
            emo_counts = df_target['emotion'].value_counts().reset_index()
            emo_counts.columns = ['Emotion', 'Count']
            
            fig = px.bar(
                emo_counts, x='Emotion', y='Count',
                title=f"顧客對「{selected_food}」的情緒反應分佈",
                color='Emotion',
                text_auto=True,
                color_discrete_sequence=px.colors.qualitative.Pastel
            )
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("此餐點無數據")

    st.divider()

    # ==========================================
    # UI 區塊 C: 詳細佐證與人工驗證
    # ==========================================
    st.markdown(f"### ✅ 資料驗證 ({len(df_target)} 筆)")
    st.caption("如果您發現 AI 判斷錯誤 (例如：這不是漢堡，或表情判斷錯誤)，請取消勾選，上方的統計圖表會自動扣除該筆數據。")

    # 使用 Grid 顯示
    cols = st.columns(3)
    
    for i, row in df_target.iterrows():
        col = cols[i % 3]
        with col:
            with st.container(border=True):
                # 標題
                st.markdown(f"**{row['emotion']}** <span style='color:gray'>({row['timestamp']})</span>", unsafe_allow_html=True)
                
                # 左右並排顯示圖
                img_c1, img_c2 = st.columns(2)
                with img_c1:
                    if os.path.exists(row['face_path']):
                        st.image(row['face_path'], use_container_width=True)
                    else: st.text("No Face")
                with img_c2:
                    st.image(row['plate_path'], use_container_width=True)

                # 勾選框 (互動核心)
                # 當使用者改變勾選狀態時，會呼叫 db_manager 更新資料庫，然後 Streamlit 會自動重跑 (Rerun)
                checkbox_key = f"chk_food_{row['id']}"
                
                def on_change_callback(eid=row['id'], k=checkbox_key):
                    # 取得最新狀態
                    new_val = st.session_state[k]
                    # 更新資料庫
                    db_manager.update_evidence_feedback(eid, new_val)
                    # 提示
                    if not new_val:
                        st.toast(f"已移除 ID {eid}，圖表將重新計算")

                st.checkbox(
                    "資料正確 (納入統計)", 
                    value=True, 
                    key=checkbox_key,
                    on_change=on_change_callback
                )

def _render_all_emotions_gallery(db_manager, session_id):
    """
    [NEW] 顯示所有偵測到的情緒照片 (含 Top 2 分數)
    """
    df = db_manager.get_event_evidence(session_id, "strong_emotion_face")
    
    if df.empty:
        st.info("尚無情緒紀錄")
        return

    # 使用 Grid 佈局
    cols = st.columns(4)
    
    for i, row in df.iterrows():
        path = row['local_path']
        if not os.path.exists(path): continue
            
        filename = os.path.basename(path)
        # 解析檔名: 時間_情緒1-分數_情緒2-分數_Face.jpg
        try:
            parts = filename.split('_')
            # parts[0]: 日期, parts[1]: 時間
            time_str = f"{parts[1]}" 
            
            # 解析情緒 1 (例如 "開心-98")
            e1_part = parts[2].split('-')
            e1_label = e1_part[0]
            e1_score = e1_part[1]
            
            # 解析情緒 2 (例如 "驚艷-02")
            # 舊的檔案可能沒有第二情緒，要做防呆
            if len(parts) >= 5:
                e2_part = parts[3].split('-')
                e2_label = e2_part[0]
                e2_score = e2_part[1]
                caption_text = f"🥇{e1_label}({e1_score}%) | 🥈{e2_label}({e2_score}%)"
            else:
                caption_text = f"🥇{e1_label}({e1_score}%)"
                
        except:
            # 解析失敗 (可能是舊檔案)
            time_str = "Unknown"
            caption_text = "Legacy Data"

        col = cols[i % 4]
        with col:
            st.image(path, use_container_width=True)
            st.caption(f"🕒 {time_str}")
            st.markdown(f"**{caption_text}**")

# ==========================================
# [重構] 獨立的 Tab 渲染函式 (Function Components)
# ==========================================




def _render_tab_global(client, db_manager, df_sessions, t):
    """Tab 5: 菜色研發報告 (Menu R&D Report)"""
    st.subheader("🍔 菜色研發報告 (Menu Insights)")
    st.caption("針對特定菜色的顧客情緒反應進行分析，適合主廚與菜單研發人員。")

    if df_sessions.empty:
        st.warning("⚠️ 目前選定的時間範圍內無 Session 資料，無法分析菜色。")
        return
    
    # --- 1. 資料聚合 (Aggregation) ---
    # 這裡的邏輯是將所有場次的「餐點情緒」彙整起來
    all_food_data = []
    
    for _, session_row in df_sessions.iterrows():
        sid = session_row['session_id_raw']
        # 撈取該 Session 的強烈情緒餐點證據
        df_evidence = db_manager.get_event_evidence(sid, "strong_emotion_plate")
        
        if df_evidence.empty: continue
            
        for _, row in df_evidence.iterrows():
            if row['human_corrected'] == 0: continue # 排除人工否決的
            
            food_name = row['food_label'] if row['food_label'] else "Unknown"
            
            # 解析情緒 (從檔名: 日期_情緒-分數_...)
            try:
                fname = os.path.basename(row['local_path'])
                parts = fname.split('_')
                # 假設檔名結構固定，取第3部分的情緒標籤
                # 範例: ..._開心-98_...
                emotion_tag = parts[2].split('-')[0] 
                
                all_food_data.append({
                    "food": food_name,
                    "emotion": emotion_tag
                })
            except:
                continue

    if not all_food_data:
        st.info("在此時間範圍內，尚未蒐集到足夠的菜色情緒樣本 (需觸發強烈情緒快照)。")
        return

    # 轉換為 DataFrame 方便統計
    df_analysis = pd.DataFrame(all_food_data)

    # 準備給 LLM 的統計數據： {'漢堡': {'開心': 5, '嫌棄': 1}, ...}
    food_stats = {}
    for food in df_analysis['food'].unique():
        sub_df = df_analysis[df_analysis['food'] == food]
        counts = sub_df['emotion'].value_counts().to_dict()
        food_stats[food] = counts

    if "menu_report_content" not in st.session_state:
        st.session_state.menu_report_content = None

    # --- 2. AI 菜色報告區塊 ---
    with st.container(border=True):
        c1, c2 = st.columns([3, 1])
        with c1:
            st.markdown(f"**已蒐集樣本數**: `{len(df_analysis)}` 筆反應 | **涵蓋菜色**: `{len(food_stats)}` 道")
        with c2:
            gen_menu_btn = st.button("✨ 生成研發報告", type="primary", use_container_width=True)

        if gen_menu_btn:
            if not client:
                st.error("未設定 OpenAI API Key")
            else:
                with st.spinner("AI 正在分析菜色表現..."):
                    # 呼叫我們剛新增的專用函式
                    async def run_menu_gpt():
                        try:
                            resp, _ = await llm.generate_menu_report(food_stats, client)
                            return resp
                        except Exception as e:
                            return f"Error: {e}"

                    # ★ 存入 session_state
                    st.session_state.menu_report_content = asyncio.run(run_menu_gpt())
        
        # [新增] 顯示報告與下載按鈕
        if st.session_state.menu_report_content:
            st.markdown("---")
            st.markdown(st.session_state.menu_report_content)
            
            # 製作 Word 檔
            docx_file = _create_docx(st.session_state.menu_report_content)
            
            st.download_button(
                label="📥 下載菜色報告 (.docx)",
                data=docx_file,
                file_name=f"Menu_Report_{datetime.date.today()}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="secondary"
            )

    # --- 3. 單品項詳細圖表 (原本的功能) ---
    st.divider()
    st.markdown("#### 🔎 單品項詳細數據")
    
    c1, c2 = st.columns([1, 2])
    with c1:
        food_list = sorted(list(food_stats.keys()))
        selected_food = st.selectbox("選擇要鑽研的菜色", food_list)
        if selected_food:
            st.json(food_stats[selected_food])

    with c2:
        if selected_food:
            df_target = df_analysis[df_analysis['food'] == selected_food]
            emo_counts = df_target['emotion'].value_counts().reset_index()
            emo_counts.columns = ['Emotion', 'Count']
            
            fig = px.bar(
                emo_counts, x='Emotion', y='Count',
                title=f"「{selected_food}」情緒分佈圖",
                color='Emotion', text_auto=True,
                color_discrete_sequence=px.colors.qualitative.Pastel
            )
            st.plotly_chart(fig, use_container_width=True)


def _render_tab_overview(client, df_logs, num_groups, groups_df, df_sessions, stats, date_range_strs, t):
    """
    [NEW] 營運數據概觀 Tab
    整合了：關鍵指標 (KPIs)、圖表 (人流 & 情緒)、以及營運報告生成按鈕。
    """
    # --- 1. 頂部關鍵數據 (Key Metrics) ---
    st.subheader("關鍵營運指標 (Key Performance Indicators)")
    
    # 第一排：人流相關
    c1, c2, c3 = st.columns(3)
    avg_ppl = 0
    if not df_logs.empty:
        valid_ppl = df_logs[df_logs['people_count'] > 0]['people_count']
        if not valid_ppl.empty:
            avg_ppl = valid_ppl.mean()

    c1.metric(t("metric_groups"), f"{num_groups}")      # 總客組數
    c2.metric(t("metric_avg_size"), f"{avg_ppl:.1f}")   # 平均單組人數
    c3.metric(t("metric_sessions"), len(df_sessions))   # 分析場次
    
    st.write("") # 增加一點垂直間距
    
    # 第二排：滿意度與剩食 (從原本的 Tab 2 & 3 移過來)
    k1, k2, k3 = st.columns(3)
    k1.metric(t("metric_nods"), int(stats['total_nods']))
    k2.metric(t("metric_shakes"), int(stats['total_shakes']))
    k3.metric(t("metric_waste"), f"{stats['waste_rate']:.1f}%")

    st.divider()

    # --- 2. 圖表視覺化區 (Charts) ---
    chart_c1, chart_c2 = st.columns(2)
    
    # 左側：人流趨勢圖
    with chart_c1:
        st.markdown(f"#### {t('chart_traffic')}")
        with st.container(border=True):
            if not df_logs.empty:
                df_chart = df_logs.copy()
                df_chart['timestamp'] = pd.to_datetime(df_chart['timestamp'])
                df_chart = df_chart.set_index('timestamp')
                flow_data = df_chart['people_count'].resample('5T').max().fillna(0)
                st.area_chart(flow_data, color="#06b6d4", use_container_width=True)
            else:
                st.info("NO TRAFFIC DATA")
        
        # 如果有詳細組別數據，顯示在摺疊選單中
        if num_groups > 0:
            with st.expander("查看人流詳細數據 (Groups Detail)"):
                st.dataframe(groups_df, use_container_width=True, hide_index=True)
        
    # 右側：情緒分佈圖 (從原本的滿意度分析移過來)
    with chart_c2:
        st.markdown("#### 😊 情緒分佈 (Emotion Distribution)")
        
        # 篩選情緒數據
        df_emotions = df_logs[df_logs['source_type'].isin(['live_session_summary', 'uploaded_video', 'live_dual_cam'])]

        with st.container(border=True):
            if not df_emotions.empty:
                all_emotions = Counter()
                data_found = False

                for _, row in df_emotions.iterrows():
                    e_raw = row.get('emotions')
                    if pd.isna(e_raw) or e_raw == "": continue
                    try:
                        # 處理字串轉字典
                        e_dict = ast.literal_eval(str(e_raw)) if isinstance(e_raw, str) else e_raw
                        if isinstance(e_dict, dict):
                            for k, v in e_dict.items():
                                if k not in ['Meal_Status', 'status']:
                                    try:
                                        val = float(v) 
                                        if val > 0:
                                            all_emotions[k] += val
                                            data_found = True
                                    except: continue
                    except: continue 

                if data_found and all_emotions:
                    e_df = pd.DataFrame(all_emotions.items(), columns=['Emotion', 'Count'])
                    fig = px.bar(
                        e_df, x='Emotion', y='Count', 
                        color_discrete_sequence=['#8b5cf6'], text_auto=True
                    )
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    st.info("尚無詳細情緒數據")
            else:
                st.info("NO DATA")

    st.divider()

    # --- 3. 營運報告生成區 (Report Generator) ---
    # 把原本 Tab 4 的上半部按鈕移到這裡
    st.subheader("✨ 智慧營運顧問")
    # 初始化 session state 用來存報告
    if "op_report_content" not in st.session_state:
        st.session_state.op_report_content = None
    
    start_dt_str, end_dt_str = date_range_strs
    
    with st.container(border=True):
        c1, c2 = st.columns([3, 1])
        with c1:
            st.markdown(f"**分析區間**: `{start_dt_str}` ~ `{end_dt_str}`")
            st.info("點擊右側按鈕，讓 AI 為您總結本時段的人流、滿意度與營運狀況，並提供行動建議。")
        
        with c2:
            gen_btn = st.button("生成營運總結報告", type="primary", use_container_width=True)

        if gen_btn:
            # 準備數據
            traffic_trend_str = "數據處理中..."
            # (簡化的流量字串處理邏輯，保持原樣或略過細節以節省篇幅)
            
            # 真實人數計算
            real_total_customers = int(groups_df['最大人數'].sum()) if (not groups_df.empty and '最大人數' in groups_df.columns) else 0
            
            op_stats = {
                "total_customers": real_total_customers,
                "total_sessions": len(df_sessions),
                "satisfaction_index": f"{stats['total_nods']} (Pos) vs {stats['total_shakes']} (Neg)",
                "waste_rate": f"{stats['waste_rate']:.1f}%",
                "traffic_trend": "詳見圖表"
            }

            with st.spinner("AI 顧問正在分析營運數據..."):
                prompt = f"Analyze Operation Stats: {op_stats}"
                async def run_op_rep():
                    try: 
                        # 呼叫後端 LLM
                        BACKEND_CONFIG = {"store_type": "Buffet", "tone": "專業客觀", "tips_style": "營運流程優化"}
                        resp, _ = await llm.summarize_session(op_stats, client=client, custom_instructions=prompt, **BACKEND_CONFIG)
                        return resp
                    except Exception as e:
                        return f"Error: {e}"
                

                st.session_state.op_report_content = asyncio.run(run_op_rep())

        # [新增] 顯示報告與下載按鈕 (只要 session_state 有資料就顯示)
        if st.session_state.op_report_content:
            st.markdown("---")
            st.markdown(st.session_state.op_report_content)
            
            # 製作 Word 檔
            docx_file = _create_docx(st.session_state.op_report_content)
            
            st.download_button(
                label="📥 下載 Word 報告 (.docx)",
                data=docx_file,
                file_name=f"Operational_Report_{datetime.date.today()}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="secondary"
            )

# AI Agent
def _render_tab_ai_agent(client, db_manager, df_sessions, df_logs, stats, t):
    """
    [NEW] AI Agent 智慧對話 Tab (最終完整版)
    包含：UX 優化、RAG 資料注入、Text-to-SQL 雙階段推理、資料庫欄位自動適配
    """
    import pandas as pd
    import asyncio
    import os
    import sqlite3
    import re

    # --- 1. CSS 美化注入 (霓虹暗黑風格) ---
    st.markdown("""
    <style>
        /* 聊天視窗容器調整 */
        .stChatContainer { padding-right: 10px; }
        
        /* 1. 對話外框容器美化 */
        [data-testid="stVerticalBlockBorderWrapper"] > div {
            border-radius: 15px !important;
            border: 1px solid rgba(255, 255, 255, 0.1) !important;
            background-color: #1e293b !important;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1), 0 2px 4px -1px rgba(0, 0, 0, 0.06) !important;
        }

        /* 2. 對話氣泡美化 */
        .stChatMessage {
            background-color: transparent !important;
            padding: 1rem !important;
            border-radius: 12px !important;
            margin-bottom: 0.5rem !important;
        }

        /* AI (Assistant) - 亮青色風格 */
        .stChatMessage[data-testid="stChatMessage"]:nth-child(odd) {
            background-color: rgba(6, 182, 212, 0.1) !important;
            border-left: 3px solid #06b6d4 !important;
        }

        /* User - 淡灰色風格 */
        .stChatMessage[data-testid="stChatMessage"]:nth-child(even) {
            background-color: rgba(255, 255, 255, 0.05) !important;
        }

        /* 3. 文字與頭像優化 */
        .stChatMessage p {
            font-size: 1.1rem !important;
            line-height: 1.6 !important;
            color: #e2e8f0 !important;
        }
        .stChatMessage .stImage {
            width: 45px !important;
            height: 45px !important;
            border-radius: 50% !important;
            border: 2px solid #334155 !important;
        }
        
        /* 4. 按鈕優化 */
        button[kind="secondary"] {
            border: 1px solid rgba(255,255,255,0.2) !important;
            background-color: transparent !important;
            color: #94a3b8 !important;
        }
        button[kind="secondary"]:hover {
            border-color: #ef4444 !important;
            color: #ef4444 !important;
            background-color: rgba(239, 68, 68, 0.1) !important;
        }
        div.stButton > button {
            border-radius: 20px !important;
            transition: all 0.3s ease;
        }
    </style>
    """, unsafe_allow_html=True)

    # --- 2. 初始化 Session State ---
    if "messages" not in st.session_state:
        st.session_state.messages = [
            {"role": "assistant", "content": "👋 嗨！我是水母哥，您的智能小助手。有什麼問題都可以問我呦！"}
        ]

    # --- 3. 準備 Context (AI 的大腦) ---
    
    # (A) 計算熱門時段
    peak_hour = "資料不足"
    if not df_logs.empty:
        try:
            df_logs['hour'] = pd.to_datetime(df_logs['timestamp']).dt.hour
            peak_hour = f"{df_logs['hour'].mode()[0]}點"
        except: pass

    # (B) 撈取食物數據
    food_summary_list = []
    if not df_sessions.empty:
        # 限制前 50 筆以優化效能
        for _, row in df_sessions.head(50).iterrows():
            sid = row.get('session_id_raw')
            if not sid: continue
            
            s_time = row['timestamp'].strftime('%H:%M')
            try:
                evidence_df = db_manager.get_event_evidence(sid, "strong_emotion_plate")
                if not evidence_df.empty:
                    for _, e_row in evidence_df.iterrows():
                        if e_row['human_corrected'] == 0: continue
                        f_label = e_row['food_label']
                        if f_label:
                            food_summary_list.append(f"[{s_time}] {f_label}")
            except: pass

    food_context_str = ", ".join(food_summary_list) if food_summary_list else "目前區間內無 AI 辨識到的餐點紀錄"
    if len(food_context_str) > 2000: food_context_str = food_context_str[:2000] + "..."

    # (C) Text-to-SQL Schema 定義 (告訴 AI 資料庫長怎樣)
    # 特別說明 waste_count 是我們稍後會手動計算生成的
    db_schema_context = """
    [資料庫權限]
    你有權限存取一個 SQLite 資料庫，包含以下兩張表：
    
    1. 表名: sessions (每一筆代表一組客人的用餐紀錄)
       - Columns: 
         - nod_count (點頭次數/int)
         - shake_count (搖頭次數/int)
         - waste_count (剩食數量/int) (若大於0代表有浪費)
         - timestamp (時間/datetime)
    
    2. 表名: logs (每一筆代表攝影機抓到的人流紀錄)
       - Columns: 
         - people_count (人數/int)
         - timestamp (時間/datetime)

    [指令]
    如果使用者問統計類問題(如平均、總和、特定時段)，請生成 SQL 查詢。
    格式要求：只輸出 `SQL_QUERY: SELECT ...`，不要包含其他文字。
    """

    summary_context = f"""
    [營運摘要]
    - 場次: {len(df_sessions)}
    - 滿意: {stats['total_nods']} | 不滿: {stats['total_shakes']}
    - 剩食數: {stats['waste_count']}
    - 高峰: {peak_hour}
    [餐點紀錄] {food_context_str}
    """

    # --- 4. 介面主體：卡片式容器 ---
    with st.container(border=True):
        
        # Header
        col_header_L, col_header_R = st.columns([5, 1])
        with col_header_L:
            c_img, c_txt = st.columns([1, 6])
            with c_img:
                img_path = "assets/avatar.png"
                if os.path.exists(img_path): st.image(img_path, width=150)
                else: st.markdown("🐙")
            with c_txt:
                st.markdown("### 💬 智能小助手 - 水母哥")
                st.caption("24小時 AI 營運顧問 | 支援 SQL 數據查詢")
        with col_header_R:
            if st.button("🗑 清空", type="secondary", use_container_width=True):
                st.session_state.messages = [{"role": "assistant", "content": "紀錄已清空！"}]
                st.rerun()

        st.divider()

        # 對話捲動區塊
        chat_container = st.container(height=400)
        with chat_container:
            for msg in st.session_state.messages:
                role = msg["role"]
                # 設定頭像
                if role == "user":
                    avatar = "👤"
                else:
                    avatar = "assets/avatar.png" if os.path.exists("assets/avatar.png") else "🤖"
                
                with st.chat_message(role, avatar=avatar):
                    st.markdown(msg["content"])
                    # # 如果有 SQL 執行結果，顯示在摺疊選單中
                    # if "sql_query" in msg:
                    #     st.caption(f"🔍 SQL: `{msg['sql_query']}`")
                    #     with st.expander("查看原始數據"):
                    #         st.code(msg.get('sql_result', 'No Data'))

    # --- 5. 快捷按鈕區 ---
    st.write("💡 **快捷提問：**")
    b1, b2, b3, b4 = st.columns(4)
    user_click_prompt = None

    if b1.button("📊 今日總結"): user_click_prompt = "請總結今天的營運狀況與關鍵數據。"
    if b2.button("🍔 熱門餐點"): user_click_prompt = "大家都點了什麼？有沒有特定時段偏好？"
    if b3.button("📈 平均滿意度"): user_click_prompt = "平均每組客人的滿意點頭次數是多少？"
    if b4.button("🗑️ 剩食分析"): user_click_prompt = "總共有多少場次出現剩食？比例是多少？"

    # --- 6. 輸入處理邏輯 ---
    chat_input_text = st.chat_input("輸入問題...")
    final_prompt = user_click_prompt if user_click_prompt else chat_input_text

    if final_prompt:
        # 1. 顯示使用者訊息
        st.session_state.messages.append({"role": "user", "content": final_prompt})
        with chat_container:
            with st.chat_message("user", avatar="👤"):
                st.markdown(final_prompt)

        # 2. AI 處理 (Text-to-SQL Magic)
        if not client:
            st.error("⚠️ 未設定 OpenAI API Key")
        else:
            with chat_container:
                avatar = "assets/avatar.png" if os.path.exists("assets/avatar.png") else "🤖"
                with st.chat_message("assistant", avatar=avatar):
                    status_placeholder = st.empty()
                    
                    with st.spinner("水母哥正在思考..."):
                        async def run_analysis():
                            # System Prompt 包含 Schema
                            full_prompt = f"""
                            你是一位專業餐廳顧問。
                            {summary_context}
                            {db_schema_context}
                            請根據使用者問題判斷：
                            1. 若是閒聊或摘要，直接回答。
                            2. 若需計算(平均/加總/過濾)，請生成 `SQL_QUERY: SELECT ...`。
                            3.【注意】請直接說出結論或數字即可，完全不要提到「SQL」、「資料庫」或「查詢語句」等技術字眼。語氣要像是一位專業的店長在做匯報。
                            """
                            
                            # A. 第一次請求
                            resp = await client.chat.completions.create(
                                model="gpt-4o",
                                messages=[{"role": "system", "content": full_prompt}] + st.session_state.messages,
                                temperature=0
                            )
                            first_reply = resp.choices[0].message.content
                            
                            # B. 檢查 SQL
                            sql_match = re.search(r"SQL_QUERY:\s*(SELECT.*)", first_reply, re.IGNORECASE | re.DOTALL)
                            
                            if sql_match:
                                sql_query = sql_match.group(1).strip().replace("```sql", "").replace("```", "").strip()
                                status_placeholder.markdown(f"⚡️ 水母哥正在查詢資料庫...")
                                
                                # C. 建立內存資料庫 (解決 leftover_data 問題)
                                try:
                                    conn = sqlite3.connect(':memory:')
                                    
                                    # --- 處理 Sessions 表 ---
                                    clean_sessions = df_sessions.copy()
                                    # [關鍵邏輯] 將 leftover_data (JSON字串) 轉為 waste_count (Int)
                                    if 'leftover_data' in clean_sessions.columns:
                                        clean_sessions['waste_count'] = clean_sessions['leftover_data'].apply(
                                            lambda x: 1 if x and isinstance(x, str) and len(x) > 4 else 0
                                        )
                                    else:
                                        clean_sessions['waste_count'] = 0
                                    
                                    # 補齊欄位
                                    for col in ['nod_count', 'shake_count', 'timestamp']:
                                        if col not in clean_sessions.columns: clean_sessions[col] = 0
                                    
                                    clean_sessions = clean_sessions[['nod_count', 'shake_count', 'waste_count', 'timestamp']].fillna(0)
                                    clean_sessions.to_sql('sessions', conn, index=False)
                                    
                                    # --- 處理 Logs 表 ---
                                    clean_logs = df_logs.copy()
                                    clean_logs = clean_logs[['people_count', 'timestamp']].fillna(0)
                                    clean_logs.to_sql('logs', conn, index=False)
                                    
                                    # 執行 SQL
                                    query_df = pd.read_sql_query(sql_query, conn)
                                    result_str = query_df.to_string()
                                    conn.close()
                                    
                                    # D. 第二次請求 (解釋結果)
                                    final_prompt_sys = f"SQL查詢: {sql_query}\n結果:\n{result_str}\n請根據結果用繁體中文回答。"
                                    resp2 = await client.chat.completions.create(
                                        model="gpt-4o",
                                        messages=[{"role": "system", "content": final_prompt_sys}],
                                        temperature=0.7
                                    )
                                    return resp2.choices[0].message.content, sql_query, result_str
                                    
                                except Exception as e:
                                    return f"查詢失敗: {e}", None, None
                            else:
                                return first_reply, None, None

                        reply_text, executed_sql, sql_result = asyncio.run(run_analysis())
                        
                        status_placeholder.empty()
                        
                        # 儲存與顯示
                        msg_data = {"role": "assistant", "content": reply_text}
                        if executed_sql:
                            msg_data["sql_query"] = executed_sql
                            msg_data["sql_result"] = sql_result
                        
                        st.session_state.messages.append(msg_data)
                        
                        # 顯示這次的回答 (因為 rerun 會清掉畫面，所以存檔後直接 rerun 讓迴圈顯示)
                        # 但為了避免瞬間空白，我們可以選擇這裡不 render，直接交給 rerun
        
        # 3. 強制重整 (確保流暢)
        st.rerun()

def _render_tab_evidence(db, df_sessions, t):
    """
    [NEW] 影像佐證 Tab
    專門顯示每一個 Session 的詳細照片 (Nod, Shake, Waste, etc.)
    """
    st.subheader(f"{t('header_evidence')} ({len(df_sessions)})")
    st.caption("以下列出篩選時段內的所有用餐紀錄及其影像佐證。")
    
    if not df_sessions.empty:
        # 確保有 raw_id
        if 'session_id_raw' not in df_sessions.columns:
            df_sessions['session_id_raw'] = df_sessions['timestamp'].dt.strftime('%Y%m%d%H%M%S')

        # 顯示列表 (Expander List)
        for _, row in df_sessions.iterrows():
            ts = row['timestamp']
            time_str = ts.strftime('%m/%d %H:%M')
            unique_session_id = row['session_id_raw']
            
            nods = int(row.get('nod_count', 0))
            shakes = int(row.get('shake_count', 0))
            
            # 標題顯示時間與簡易情緒統計
            label = f"📍 {time_str} | 😊 {nods} vs 😟 {shakes}"

            with st.expander(label, expanded=False):
                # 這裡保留原本的詳細 Tabs
                t1, t2, t3, t4, t5 = st.tabs(["🎥 點頭 (Nod)", "🎥 搖頭 (Shake)", "🍽️ 剩食 (Waste)", "🔥 交叉比對", "😊 情緒快照"])
                
                with t1: _render_evidence_grid(db, unique_session_id, 'nod')
                with t2: _render_evidence_grid(db, unique_session_id, 'shake')
                with t3: _render_evidence_grid(db, unique_session_id, 'plate_vlm')
                with t4: _render_comparison_gallery(db, unique_session_id)
                with t5: _render_all_emotions_gallery(db, unique_session_id)
    else:
        st.info("在此區間內無資料。")

# ==========================================
# 主顯示函式 (Controller)
# ==========================================

def display(client, db_manager, t=None): 
    if t is None: 
        def t(k): return k
    
    db = db_manager 

    col_title, col_refresh = st.columns([5, 1])
    with col_title:
        st.subheader(t("dash_title"))
    with col_refresh:
        if st.button(t("btn_refresh"), use_container_width=True):
            st.rerun()

    # 1. 篩選器 (Filter Section)
    with st.container(border=True):
        st.markdown(f"<h5 style='color:var(--primary-color); font-weight:bold;'>{t('filter_title')}</h5>", unsafe_allow_html=True)
        col1, col2, col3 = st.columns(3)
        with col1:
            today = datetime.date.today()
            date_range = st.date_input(t("date_range"), value=[today, today], format="YYYY/MM/DD")
        with col2:
            time_range_option = st.selectbox(t("time_period"), [t("opt_all_day"), t("opt_custom")])
        with col3:
            source_option = st.selectbox(t("data_source"), ["All", "Live", "Video"])

        if len(date_range) != 2:
            st.warning("Please select end date.")
            st.stop()

        start_date, end_date = date_range
        if "All Day" in time_range_option or "全日" in time_range_option:
            start_dt_str = f"{start_date} 00:00:00"
            end_dt_str = f"{end_date} 23:59:59"
        else:
            col_start, col_end = st.columns(2)
            with col_start:
                s_time = st.time_input("Start", datetime.time(9, 0))
            with col_end:
                e_time = st.time_input("End", datetime.time(21, 0))
            start_dt_str = f"{start_date} {s_time.strftime('%H:%M:%S')}"
            end_dt_str = f"{end_date} {e_time.strftime('%H:%M:%S')}"
            
    # 2. 數據獲取 (Data Fetching)
    if source_option == "Live":
        selected_sources = ['live_stream', 'live_session_summary', 'live_dual_cam']
    elif source_option == "Video":
        selected_sources = ['uploaded_video']
    else:
        selected_sources = ['live_stream', 'live_session_summary', 'uploaded_video', 'live_dual_cam']

    df_logs = db.get_logs_by_range(start_dt_str, end_dt_str, source_types=selected_sources)
    num_groups, groups_df = db.get_customer_groups_analysis(start_dt_str, end_dt_str, gap_minutes=0.6)
    df_sessions_all = db.get_all_session_records()
    
    df_sessions = pd.DataFrame()
    if not df_sessions_all.empty:
        df_sessions_all['timestamp'] = pd.to_datetime(df_sessions_all['timestamp'])
        mask = (df_sessions_all['timestamp'] >= pd.to_datetime(start_dt_str)) & \
               (df_sessions_all['timestamp'] <= pd.to_datetime(end_dt_str))
        df_sessions = df_sessions_all.loc[mask].copy()
        df_sessions = df_sessions.sort_values('timestamp', ascending=False)

    if df_logs.empty and df_sessions.empty and num_groups == 0:
        st.info("NO DATA AVAILABLE.")
        return

    # 3. 預先計算共用統計數據 (Pre-calculate Stats)
    total_nods = df_sessions['nod_count'].sum() if not df_sessions.empty else 0
    total_shakes = df_sessions['shake_count'].sum() if not df_sessions.empty else 0
    waste_count = 0
    if not df_sessions.empty:
        for _, row in df_sessions.iterrows():
            try:
                data = json.loads(row['leftover_data'])
                if data and len(data) > 0: 
                    waste_count += 1
            except: pass
    waste_rate = (waste_count / len(df_sessions) * 100) if not df_sessions.empty else 0
    
    stats = {
        'total_nods': total_nods,
        'total_shakes': total_shakes,
        'waste_count': waste_count,
        'waste_rate': waste_rate
    }

    # 4. 主分頁顯示 (Main Tabs)
    tab1, tab2, tab3, tab4 = st.tabs([
        t("tab_overview"),      # 📊 營運數據概觀
        t("tab_menu_insight"),  # 🍔 菜色整體洞察
        t("tab_ai_agent"),     # 🤖 AI Agent 智慧洞察
        t("tab_evidence")      # 📸 區間影像佐證紀錄
    ])

    # Tab 1: 營運數據概觀 (合併了人流、滿意度、圖表、報告按鈕)
    with tab1:
        _render_tab_overview(
            client, 
            df_logs, 
            num_groups, 
            groups_df, 
            df_sessions, 
            stats, 
            (start_dt_str, end_dt_str), 
            t
        )

    # Tab 3: 菜色整體洞察 (原本的 Global Insight，邏輯不變，只是換位置)
    with tab2:
        _render_tab_global(client, db, df_sessions, t)

    # Tab 4: AI Agent (目前留空)
    with tab3:
        _render_tab_ai_agent(client, db, df_sessions, df_logs, stats, t)
        # with st.container(border=True):
        #     st.info("🚧 **AI Agent 智慧洞察功能開發中**")
        #     st.markdown("""
        #     未來功能預告：
        #     - 🗣️ **自然語言對話**：直接問系統「上週五中午生意好嗎？」
        #     - 🤖 **自動化任務**：設定條件自動發送 Line 通知。
        #     - 🧠 **深度關聯分析**：分析天氣、促銷活動與情緒的關聯。
        #     """)

        # Tab 2: 影像佐證紀錄 (獨立出來的照片區)
    with tab4:
        _render_tab_evidence(db, df_sessions, t)